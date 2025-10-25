"""
파일 파싱 유틸리티 - TXT, PDF, DOCX 파일에서 텍스트 추출
"""

import logging
from pathlib import Path
from typing import Tuple, Optional
import chardet

logger = logging.getLogger(__name__)

# 파일 설정
SUPPORTED_FORMATS = {'.txt', '.pdf', '.docx'}
MAX_FILE_SIZE_MB = 50
MAX_FILE_SIZE_BYTES = MAX_FILE_SIZE_MB * 1024 * 1024


class FileParser:
    """다양한 형식의 파일에서 텍스트를 추출하는 클래스"""

    @staticmethod
    def parse(file_path: str) -> Tuple[str, str]:
        """
        파일을 파싱하여 텍스트 추출

        Args:
            file_path: 파일 경로

        Returns:
            (text: 추출된 텍스트, format: 파일 형식)

        Raises:
            ValueError: 지원하지 않는 형식 또는 크기 초과
            IOError: 파일 읽기 실패
        """
        path = Path(file_path)

        # 파일 존재 확인
        if not path.exists():
            raise IOError(f"파일을 찾을 수 없음: {file_path}")

        # 파일 크기 확인
        file_size = path.stat().st_size
        if file_size > MAX_FILE_SIZE_BYTES:
            raise ValueError(f"파일 크기 초과: {file_size / 1024 / 1024:.1f}MB (최대: {MAX_FILE_SIZE_MB}MB)")

        # 파일 형식 확인
        suffix = path.suffix.lower()
        if suffix not in SUPPORTED_FORMATS:
            raise ValueError(f"지원하지 않는 파일 형식: {suffix}. 지원: {', '.join(SUPPORTED_FORMATS)}")

        # 파일 형식별 파싱
        try:
            if suffix == '.txt':
                text = FileParser._parse_txt(file_path)
            elif suffix == '.pdf':
                text = FileParser._parse_pdf(file_path)
            elif suffix == '.docx':
                text = FileParser._parse_docx(file_path)
            else:
                raise ValueError(f"알 수 없는 파일 형식: {suffix}")

            logger.info(f"✅ 파일 파싱 완료: {path.name} ({len(text)}자)")
            return text, suffix[1:]  # 점 제거

        except Exception as e:
            logger.error(f"❌ 파일 파싱 실패: {path.name} - {str(e)}")
            raise

    @staticmethod
    def _parse_txt(file_path: str) -> str:
        """TXT 파일 파싱"""
        try:
            # 인코딩 자동 감지
            with open(file_path, 'rb') as f:
                raw = f.read()

            detected = chardet.detect(raw)
            encoding = detected.get('encoding', 'utf-8')

            if encoding is None:
                encoding = 'utf-8'

            try:
                text = raw.decode(encoding)
            except (UnicodeDecodeError, LookupError):
                # 인코딩 감지 실패 시 utf-8, cp949 시도
                for enc in ['utf-8', 'cp949', 'euc-kr', 'latin-1']:
                    try:
                        text = raw.decode(enc)
                        break
                    except (UnicodeDecodeError, LookupError):
                        continue
                else:
                    raise ValueError("텍스트 인코딩을 자동 감지할 수 없습니다")

            return text.strip()

        except Exception as e:
            raise IOError(f"TXT 파일 읽기 실패: {str(e)}")

    @staticmethod
    def _parse_pdf(file_path: str) -> str:
        """PDF 파일 파싱"""
        try:
            import PyPDF2

            text_parts = []
            with open(file_path, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)

                if len(pdf_reader.pages) == 0:
                    raise ValueError("PDF 파일에 페이지가 없습니다")

                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        text = page.extract_text()
                        if text:
                            text_parts.append(text)
                    except Exception as e:
                        logger.warning(f"PDF 페이지 {page_num} 파싱 실패: {str(e)}")
                        continue

            if not text_parts:
                raise ValueError("PDF에서 텍스트를 추출할 수 없습니다")

            return '\n'.join(text_parts).strip()

        except ImportError:
            raise IOError("PyPDF2 라이브러리가 설치되지 않았습니다. pip install PyPDF2를 실행하세요.")
        except Exception as e:
            raise IOError(f"PDF 파일 파싱 실패: {str(e)}")

    @staticmethod
    def _parse_docx(file_path: str) -> str:
        """DOCX 파일 파싱"""
        try:
            from docx import Document

            doc = Document(file_path)
            text_parts = []

            # 단락(paragraphs) 추출
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)

            # 테이블 추출
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' '.join(
                        cell.text.strip() for cell in row.cells if cell.text.strip()
                    )
                    if row_text:
                        text_parts.append(row_text)

            if not text_parts:
                raise ValueError("DOCX에서 텍스트를 추출할 수 없습니다")

            return '\n'.join(text_parts).strip()

        except ImportError:
            raise IOError("python-docx 라이브러리가 설치되지 않았습니다. pip install python-docx를 실행하세요.")
        except Exception as e:
            raise IOError(f"DOCX 파일 파싱 실패: {str(e)}")

    @staticmethod
    def validate_file(file_path: str) -> bool:
        """파일 유효성 검증"""
        try:
            path = Path(file_path)

            # 존재 확인
            if not path.exists():
                return False

            # 크기 확인
            if path.stat().st_size > MAX_FILE_SIZE_BYTES:
                return False

            # 형식 확인
            if path.suffix.lower() not in SUPPORTED_FORMATS:
                return False

            return True

        except Exception:
            return False

    @staticmethod
    def get_file_info(file_path: str) -> dict:
        """파일 정보 조회"""
        path = Path(file_path)
        return {
            "filename": path.name,
            "format": path.suffix[1:].lower(),
            "size_mb": round(path.stat().st_size / 1024 / 1024, 2),
            "exists": path.exists(),
        }
